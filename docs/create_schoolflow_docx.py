from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "SchoolFlow_system_analysis_design.docx"
DIAGRAM_PATH = OUT_DIR / "SchoolFlow_class_diagram.png"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        write_cell(table.rows[0].cells[index], header, True)
        shade(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            write_cell(cells[index], value)
    doc.add_paragraph()


def draw_class_diagram():
    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arial.ttf", 34)
        box_font = ImageFont.truetype("arial.ttf", 25)
        small_font = ImageFont.truetype("arial.ttf", 21)
    except OSError:
        title_font = box_font = small_font = ImageFont.load_default()

    def box(x, y, w, h, title, lines, fill):
        draw.rounded_rectangle([x, y, x + w, y + h], radius=18, fill=fill, outline=(44, 62, 80), width=3)
        draw.rectangle([x, y, x + w, y + 58], fill=(38, 99, 235), outline=(38, 99, 235))
        draw.text((x + 20, y + 15), title, fill="white", font=box_font)
        yy = y + 78
        for line in lines:
            draw.text((x + 20, yy), line, fill=(23, 33, 43), font=small_font)
            yy += 33

    draw.text((width // 2 - 330, 35), "Диаграмма классов SchoolFlow", fill=(23, 33, 43), font=title_font)
    box(80, 150, 520, 360, "Student", [
        "+ Id: int", "+ FullName: string", "+ ClassName: string", "+ Email: string?",
        "+ Login: string?", "+ PasswordHash: string?", "+ CreatedAt: DateTime",
        "+ Tasks: ICollection<SchoolTask>"
    ], (245, 248, 251))
    box(720, 150, 560, 430, "SchoolTask", [
        "+ Id: int", "+ Title: string", "+ Description: string?", "+ Subject: string",
        "+ Teacher: string?", "+ DueDate: DateTime", "+ Priority: TaskPriority",
        "+ Status: SchoolTaskStatus", "+ AssignedStudentId: int?",
        "+ AssignedStudent: Student?", "+ CreatedAt: DateTime"
    ], (245, 248, 251))
    box(80, 640, 520, 250, "TeacherAccount", [
        "+ Id: int", "+ Login: string", "+ DisplayName: string",
        "+ PasswordHash: string", "+ CreatedAt: DateTime"
    ], (245, 248, 251))
    box(1360, 150, 360, 180, "TaskPriority", ["Low = 0", "Medium = 1", "High = 2"], (250, 250, 240))
    box(1360, 400, 360, 210, "SchoolTaskStatus", ["Planned = 0", "InProgress = 1", "Completed = 2"], (250, 250, 240))
    box(720, 680, 560, 200, "SchoolTaskContext", [
        "DbSet<SchoolTask> SchoolTasks", "DbSet<Student> Students",
        "DbSet<TeacherAccount> TeacherAccounts", "Student 1..* SchoolTask"
    ], (238, 251, 246))

    draw.line([600, 300, 720, 300], fill=(15, 118, 110), width=5)
    draw.polygon([(720, 300), (695, 288), (695, 312)], fill=(15, 118, 110))
    draw.text((622, 255), "1", fill=(15, 118, 110), font=small_font)
    draw.text((675, 255), "0..*", fill=(15, 118, 110), font=small_font)
    draw.text((610, 325), "AssignedStudent", fill=(15, 118, 110), font=small_font)
    draw.line([1280, 250, 1360, 240], fill=(120, 80, 20), width=3)
    draw.line([1280, 455, 1360, 500], fill=(120, 80, 20), width=3)
    draw.line([930, 680, 930, 580], fill=(80, 80, 80), width=3)
    draw.line([720, 760, 600, 760], fill=(80, 80, 80), width=3)
    draw.line([1000, 680, 1260, 580], fill=(80, 80, 80), width=3)
    image.save(DIAGRAM_PATH)


def build_doc():
    draw_class_diagram()
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SchoolFlow")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Системный анализ и системное проектирование")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph(
        "Документ описывает веб-приложение SchoolFlow: предметную область, требования, роли пользователей, "
        "архитектуру MVC, структуру базы данных, сущности и связи между ними."
    )
    doc.add_page_break()

    doc.add_heading("Глава 1. Системный анализ", level=1)
    doc.add_heading("1.1. Характеристика предметной области", level=2)
    doc.add_paragraph(
        "SchoolFlow предназначен для автоматизации работы с учебными задачами в школе. "
        "Система объединяет новости, список класса, выдачу задач, кабинет учителя и кабинет ученика."
    )
    doc.add_paragraph(
        "Цель проекта - повысить прозрачность учебного процесса: учитель управляет заданиями и классом, "
        "а ученик получает доступ к персональному списку задач."
    )

    doc.add_heading("1.2. Цели разработки", level=2)
    for item in [
        "создать единый веб-интерфейс для школьных новостей, задач и кабинетов пользователей;",
        "обеспечить отдельный вход для учителя и ученика;",
        "дать учителю возможность импортировать список класса из Excel;",
        "позволить назначать задачи конкретным ученикам;",
        "показывать ученику только его персональные задачи;",
        "обеспечить запуск приложения на любом ПК через Docker.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.3. Пользователи системы", level=2)
    add_table(doc, ["Роль", "Описание", "Основные действия"], [
        ["Гость", "Неавторизованный посетитель", "Просмотр новостей, переход к формам входа"],
        ["Учитель", "Пользователь с правами управления учебным процессом", "Импорт класса, создание задач, аналитика"],
        ["Ученик", "Пользователь, получающий задания", "Просмотр личных задач, отметка выполнения"],
    ])

    doc.add_heading("1.4. Функциональные требования", level=2)
    add_table(doc, ["Код", "Требование"], [
        ["FR-01", "Отображение главной страницы с новостями и объявлениями."],
        ["FR-02", "Вход учителя по логину и паролю."],
        ["FR-03", "Вход ученика по логину и паролю."],
        ["FR-04", "Импорт списка учеников из Excel-файла .xlsx."],
        ["FR-05", "Создание логинов учеников при импорте."],
        ["FR-06", "Создание задачи и назначение ее конкретному ученику."],
        ["FR-07", "Редактирование, удаление и изменение статуса задачи учителем."],
        ["FR-08", "Показ ученику только назначенных ему задач."],
        ["FR-09", "Отметка учеником своей задачи выполненной."],
        ["FR-10", "Сводная аналитика для учителя."],
        ["FR-11", "Экспорт задач в CSV."],
    ])

    doc.add_heading("1.5. Нефункциональные требования", level=2)
    add_table(doc, ["Категория", "Требование"], [
        ["Переносимость", "Запуск через Docker Compose на ПК с Docker Desktop."],
        ["Удобство", "Адаптивный и понятный интерфейс для учителя и ученика."],
        ["Надежность", "Сохранение данных между перезапусками контейнера в Docker volume."],
        ["Безопасность", "Разграничение доступа ролями Teacher и Student."],
        ["Расширяемость", "Архитектура MVC должна позволять добавлять новые сущности и контроллеры."],
    ])

    doc.add_heading("1.6. Входные и выходные данные", level=2)
    add_table(doc, ["Тип данных", "Источник", "Назначение"], [
        ["Новости", "Razor-представления", "Информирование пользователей"],
        ["Список класса", "Excel-файл .xlsx", "Создание учеников и учетных данных"],
        ["Учебная задача", "Форма создания задачи", "Постановка задания ученику"],
        ["Статус задачи", "Действия учителя или ученика", "Отслеживание выполнения"],
        ["CSV-отчет", "Таблица SchoolTasks", "Выгрузка задач"],
    ])

    doc.add_page_break()
    doc.add_heading("Глава 2. Системное проектирование", level=1)
    doc.add_heading("2.1. Архитектурное решение", level=2)
    doc.add_paragraph(
        "Приложение реализовано по шаблону MVC. Контроллеры принимают HTTP-запросы, обращаются к контексту базы данных "
        "и передают подготовленные модели в Razor-представления. Модели описывают предметные сущности."
    )
    add_table(doc, ["Компонент", "Расположение", "Назначение"], [
        ["Models", "Models/", "Сущности БД, перечисления, ViewModel и формы входа"],
        ["Views", "Views/", "Razor-страницы пользовательского интерфейса"],
        ["Controllers", "Controllers/", "Маршруты и обработка сценариев пользователей"],
        ["Data", "Data/", "EF Core DbContext, инициализация БД, хеширование паролей"],
        ["wwwroot", "wwwroot/", "CSS, JavaScript и статические файлы"],
    ])

    doc.add_heading("2.2. Диаграмма классов", level=2)
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2.3. Описание классов предметной области", level=2)
    add_table(doc, ["Класс", "Назначение"], [
        ["Student", "Описывает ученика, класс, email, логин, хеш пароля и назначенные задачи."],
        ["SchoolTask", "Описывает учебную задачу, срок, предмет, статус, приоритет и ученика-адресата."],
        ["TeacherAccount", "Описывает учетную запись учителя."],
        ["SchoolTaskContext", "Контекст EF Core с DbSet для задач, учеников и учителей."],
        ["TaskPriority", "Перечисление уровней приоритета задачи."],
        ["SchoolTaskStatus", "Перечисление статусов выполнения задачи."],
    ])

    doc.add_heading("2.4. Проектирование базы данных", level=2)
    doc.add_paragraph(
        "База SQLite содержит три основные таблицы: SchoolTasks, Students и TeacherAccounts. "
        "Связь между учениками и задачами реализована через nullable-внешний ключ AssignedStudentId."
    )

    doc.add_heading("2.4.1. Таблица SchoolTasks", level=3)
    add_table(doc, ["Поле", "Тип", "Ограничения", "Описание"], [
        ["Id", "INTEGER", "PK, AUTOINCREMENT", "Идентификатор задачи"],
        ["Title", "TEXT", "NOT NULL, max 120", "Название задачи"],
        ["Description", "TEXT", "NULL, max 800", "Описание задачи"],
        ["Subject", "TEXT", "NOT NULL, max 80", "Предмет"],
        ["Teacher", "TEXT", "NULL, max 80", "Имя преподавателя"],
        ["DueDate", "TEXT", "NOT NULL", "Срок сдачи"],
        ["Priority", "INTEGER", "NOT NULL", "Приоритет: Low, Medium, High"],
        ["Status", "INTEGER", "NOT NULL", "Статус: Planned, InProgress, Completed"],
        ["AssignedStudentId", "INTEGER", "FK NULL", "Ссылка на ученика"],
        ["CreatedAt", "TEXT", "NOT NULL", "Дата создания"],
    ])

    doc.add_heading("2.4.2. Таблица Students", level=3)
    add_table(doc, ["Поле", "Тип", "Ограничения", "Описание"], [
        ["Id", "INTEGER", "PK, AUTOINCREMENT", "Идентификатор ученика"],
        ["FullName", "TEXT", "NOT NULL, max 140", "ФИО ученика"],
        ["ClassName", "TEXT", "NOT NULL, max 20", "Класс"],
        ["Email", "TEXT", "NULL, max 120", "Email ученика"],
        ["Login", "TEXT", "UNIQUE, NULL, max 80", "Логин ученика"],
        ["PasswordHash", "TEXT", "NULL, max 128", "Хеш пароля"],
        ["CreatedAt", "TEXT", "NOT NULL", "Дата добавления"],
    ])

    doc.add_heading("2.4.3. Таблица TeacherAccounts", level=3)
    add_table(doc, ["Поле", "Тип", "Ограничения", "Описание"], [
        ["Id", "INTEGER", "PK, AUTOINCREMENT", "Идентификатор учителя"],
        ["Login", "TEXT", "NOT NULL, UNIQUE, max 60", "Логин учителя"],
        ["DisplayName", "TEXT", "NOT NULL, max 140", "Отображаемое имя"],
        ["PasswordHash", "TEXT", "NOT NULL, max 128", "Хеш пароля"],
        ["CreatedAt", "TEXT", "NOT NULL", "Дата создания"],
    ])

    doc.add_heading("2.5. Связи и ограничения БД", level=2)
    add_table(doc, ["Связь / индекс", "Описание"], [
        ["Students 1 -> 0..* SchoolTasks", "Один ученик может иметь несколько назначенных задач."],
        ["SchoolTasks.AssignedStudentId -> Students.Id", "Nullable FK. Задача может быть без конкретного ученика."],
        ["IX_Students_FullName_ClassName", "Уникальный индекс для защиты от повторного импорта."],
        ["IX_Students_Login", "Уникальный индекс логина ученика."],
        ["IX_TeacherAccounts_Login", "Уникальный индекс логина учителя."],
        ["DeleteBehavior.SetNull", "При удалении ученика ссылка в задачах должна сбрасываться в NULL."],
    ])

    doc.add_heading("2.6. Проектирование авторизации", level=2)
    doc.add_paragraph(
        "В системе используется cookie-аутентификация со схемой SchoolCookie. После входа пользователю назначается роль "
        "Teacher или Student. Доступ к контроллерам ограничивается атрибутом Authorize."
    )
    add_table(doc, ["Роль", "Claims", "Доступные разделы"], [
        ["Teacher", "NameIdentifier, Name, Role=Teacher", "Кабинет учителя, класс, импорт Excel, управление задачами"],
        ["Student", "NameIdentifier, Name, Role=Student, ClassName", "Личные задачи и отметка выполнения"],
    ])

    doc.add_heading("2.7. Проектирование импорта Excel", level=2)
    doc.add_paragraph(
        "Импорт реализован в StudentsController с использованием ClosedXML. Система читает первый лист Excel-файла, "
        "пропускает строку заголовков при наличии и создает записи учеников."
    )
    add_table(doc, ["Колонка", "Обязательность", "Назначение"], [
        ["ФИО", "Да", "Полное имя ученика"],
        ["Класс", "Да", "Класс ученика"],
        ["Email", "Нет", "Контактный email"],
        ["Логин", "Нет", "Логин для входа ученика"],
        ["Пароль", "Нет", "Пароль, сохраняемый в виде хеша"],
    ])

    doc.add_heading("2.8. Развертывание", level=2)
    add_table(doc, ["Элемент", "Значение"], [
        ["Сервис", "schoolflow"],
        ["Внешний порт", "8080"],
        ["Внутренний порт", "8080"],
        ["База в контейнере", "/data/school-tasks.db"],
        ["Docker volume", "schoolflow-data"],
        ["Команда запуска", "docker compose up --build"],
    ])

    doc.add_heading("2.9. Вывод по проектированию", level=2)
    doc.add_paragraph(
        "Разработанная структура покрывает основные потребности школьного планировщика: разграничение ролей, импорт класса, "
        "назначение задач ученикам и хранение данных в SQLite. Архитектура MVC упрощает сопровождение проекта, а Docker "
        "делает запуск независимым от редактора кода."
    )

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
